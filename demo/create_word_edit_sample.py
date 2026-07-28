from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / 'demo' / 'assets' / 'word_edit_sample.docx'
BLUE = '1F4E78'
LIGHT_BLUE = 'D9EAF7'
GRAY = 'F2F2F2'


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    properties.append(shading)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn('w:tcW'))
    if width is None:
        width = OxmlElement('w:tcW')
        properties.append(width)
    width.set(qn('w:w'), str(int(width_cm * 567)))
    width.set(qn('w:type'), 'dxa')


def style_run(run, size: float = 10.5, bold: bool = False, color: str | None = None) -> None:
    run.font.name = 'Yu Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '游ゴシック')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_labeled_row(table, label: str, value: str, height_cm: float = 0.9) -> None:
    cells = table.add_row().cells
    set_cell_width(cells[0], 3.8)
    set_cell_width(cells[1], 12.7)
    set_cell_shading(cells[0], LIGHT_BLUE)
    cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
    cells[1].paragraphs[0].paragraph_format.space_after = Pt(0)
    style_run(cells[0].paragraphs[0].add_run(label), bold=True, color=BLUE)
    style_run(cells[1].paragraphs[0].add_run(value))
    row_properties = cells[0]._tc.getparent().get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    row_properties.append(cant_split)
    row_height = OxmlElement('w:trHeight')
    row_height.set(qn('w:val'), str(int(height_cm * 567)))
    row_height.set(qn('w:hRule'), 'atLeast')
    row_properties.append(row_height)


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = document.styles['Normal']
    normal.font.name = 'Yu Gothic'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '游ゴシック')
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(header.add_run('株式会社みらいワークス　経理部提出用'), 8.5, color='666666')

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    style_run(title.add_run('証 憑 紛 失 理 由 書'), 17, True, BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    style_run(subtitle.add_run('Word自動編集デモ用テンプレート'), 9, color='666666')

    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    style_run(intro.add_run('下記のとおり証憑を紛失したため、事実関係と再発防止策を報告します。'))

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = 'Table Grid'
    add_labeled_row(table, '所属', '[所属]')
    add_labeled_row(table, '氏名', '[氏名]')
    add_labeled_row(table, '紛失日', '[紛失日]')
    add_labeled_row(table, '証憑の内容', '[証憑内容]', 1.2)
    add_labeled_row(table, '金額', '[金額]')
    add_labeled_row(table, '紛失理由', '[理由]', 2.3)
    add_labeled_row(table, '再発防止策', '[再発防止策]', 2.3)

    document.add_paragraph()
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(8)
    style_run(note.add_run('確認事項'), 10, True, BLUE)
    for text in (
        '記載内容が事実と相違ないことを本人が確認してください。',
        '本システムは文書の編集を支援しますが、確定・提出・原本上書きは行いません。',
        '保存後の最終確認および社内提出は、申請者本人が行ってください。',
    ):
        paragraph = document.add_paragraph(style=None)
        paragraph.paragraph_format.left_indent = Cm(0.3)
        paragraph.paragraph_format.first_line_indent = Cm(-0.3)
        paragraph.paragraph_format.space_after = Pt(2)
        style_run(paragraph.add_run(f'・{text}'), 9)

    approval = document.add_table(rows=2, cols=3)
    approval.alignment = WD_TABLE_ALIGNMENT.RIGHT
    approval.autofit = False
    labels = ('本人確認', '上長確認', '経理確認')
    for index, label in enumerate(labels):
        set_cell_width(approval.cell(0, index), 3.3)
        set_cell_width(approval.cell(1, index), 3.3)
        set_cell_shading(approval.cell(0, index), GRAY)
        approval.cell(0, index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        approval.cell(1, index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(approval.cell(0, index).paragraphs[0].add_run(label), 8.5, True)
        style_run(approval.cell(1, index).paragraphs[0].add_run('　'), 13)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(footer.add_run('編集支援用下書き｜確定・提出は申請者本人が実施'), 8, color='777777')
    return document


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT)
    print(OUTPUT)


if __name__ == '__main__':
    main()
