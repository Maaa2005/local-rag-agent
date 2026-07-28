from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass
from typing import Any

from docx import Document


MAX_DOCX_BYTES = 10 * 1024 * 1024
MAX_UNCOMPRESSED_BYTES = 30 * 1024 * 1024
MAX_ZIP_ENTRIES = 1000


@dataclass
class EditResult:
    content: bytes
    changes: list[dict[str, Any]]
    total_replacements: int


def validate_docx_package(content: bytes) -> None:
    if not content or len(content) > MAX_DOCX_BYTES:
        raise ValueError('Wordファイルは10MB以下にしてください')
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_ZIP_ENTRIES:
                raise ValueError('Wordファイル内の項目数が上限を超えています')
            total = sum(item.file_size for item in entries)
            if total > MAX_UNCOMPRESSED_BYTES:
                raise ValueError('展開後のWordファイルが大きすぎます')
            names = {item.filename for item in entries}
            if '[Content_Types].xml' not in names or 'word/document.xml' not in names:
                raise ValueError('有効なDOCXファイルではありません')
    except zipfile.BadZipFile as exc:
        raise ValueError('有効なDOCXファイルではありません') from exc


def _replace_in_paragraph(paragraph, find: str, replacement: str) -> int:
    original = paragraph.text
    count = original.count(find)
    if count == 0:
        return 0
    if not paragraph.runs:
        paragraph.add_run(original.replace(find, replacement))
        return count

    run_ranges = []
    cursor = 0
    for index, run in enumerate(paragraph.runs):
        end = cursor + len(run.text)
        run_ranges.append((index, cursor, end))
        cursor = end

    matches = []
    start = 0
    while True:
        match_start = original.find(find, start)
        if match_start < 0:
            break
        matches.append((match_start, match_start + len(find)))
        start = match_start + len(find)

    def locate(position: int) -> tuple[int, int]:
        for index, range_start, range_end in run_ranges:
            if range_start <= position < range_end:
                return index, position - range_start
        last_index, range_start, _ = run_ranges[-1]
        return last_index, max(0, position - range_start)

    # 後方から処理すると、先行する一致位置のオフセットを変えずに書式付きrunを保てる。
    for match_start, match_end in reversed(matches):
        start_run, start_offset = locate(match_start)
        end_run, end_offset_last = locate(match_end - 1)
        end_offset = end_offset_last + 1
        if start_run == end_run:
            text = paragraph.runs[start_run].text
            paragraph.runs[start_run].text = (
                text[:start_offset] + replacement + text[end_offset:]
            )
            continue
        start_text = paragraph.runs[start_run].text
        end_text = paragraph.runs[end_run].text
        paragraph.runs[start_run].text = start_text[:start_offset] + replacement
        for index in range(start_run + 1, end_run):
            paragraph.runs[index].text = ''
        paragraph.runs[end_run].text = end_text[end_offset:]
    return count


def _iter_paragraphs(document):
    for index, paragraph in enumerate(document.paragraphs, 1):
        yield f'本文 段落{index}', paragraph
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            for cell_index, cell in enumerate(row.cells, 1):
                for paragraph_index, paragraph in enumerate(cell.paragraphs, 1):
                    location = (
                        f'表{table_index} 行{row_index} 列{cell_index} '
                        f'段落{paragraph_index}'
                    )
                    yield location, paragraph
    for section_index, section in enumerate(document.sections, 1):
        for area_name, area in [('ヘッダー', section.header), ('フッター', section.footer)]:
            for paragraph_index, paragraph in enumerate(area.paragraphs, 1):
                yield f'{area_name}{section_index} 段落{paragraph_index}', paragraph


def edit_docx(content: bytes, operations: list[dict[str, str]]) -> EditResult:
    validate_docx_package(content)
    if not 1 <= len(operations) <= 20:
        raise ValueError('編集指示は1件以上20件以下にしてください')
    document = Document(io.BytesIO(content))
    changes: list[dict[str, Any]] = []
    total = 0
    for operation_index, operation in enumerate(operations, 1):
        find = str(operation.get('find', ''))
        replacement = str(operation.get('replacement', ''))
        if not find or len(find) > 500 or len(replacement) > 2000:
            raise ValueError('検索文字は1-500文字、置換文字は2000文字以下にしてください')
        operation_count = 0
        locations = []
        for location, paragraph in _iter_paragraphs(document):
            count = _replace_in_paragraph(paragraph, find, replacement)
            if count:
                operation_count += count
                locations.append(location)
        total += operation_count
        changes.append({
            'operation': operation_index,
            'find': find,
            'replacement': replacement,
            'count': operation_count,
            'locations': locations[:20],
        })
    if total == 0:
        raise ValueError('指定された文字列は文書内に見つかりませんでした')
    output = io.BytesIO()
    document.save(output)
    return EditResult(output.getvalue(), changes, total)
