import io
import json
import asyncio

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.api import documents as documents_api
import app.database as db_module
from app.services.docx_editor import edit_docx, validate_docx_package


def make_docx() -> bytes:
    document = Document()
    document.add_heading('経費申請書', level=1)
    document.add_paragraph('申請者: 山田 太郎')
    document.add_paragraph('金額: 3,000円')
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = '利用日'
    table.cell(0, 1).text = '2026年7月1日'
    document.sections[0].header.paragraphs[0].text = '社内限定'
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def document_text(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    parts.extend(paragraph.text for paragraph in document.sections[0].header.paragraphs)
    return '\n'.join(parts)


def authenticated_headers(client) -> dict[str, str]:
    login = client.post('/api/auth/token', data={'username': 'admin', 'password': 'admin'})
    token = login.json()['access_token']
    headers = {'Authorization': 'Bearer ' + token}
    changed = client.post(
        '/api/auth/password',
        headers=headers,
        json={'current_password': 'admin', 'new_password': 'AdminPass2026!'},
    )
    assert changed.status_code == 200
    login = client.post(
        '/api/auth/token', data={'username': 'admin', 'password': 'AdminPass2026!'}
    )
    return {'Authorization': 'Bearer ' + login.json()['access_token']}


@pytest.fixture()
def client(tmp_path, monkeypatch):
    db_path = tmp_path / 'app.db'
    monkeypatch.setattr(db_module, 'DB_PATH', db_path)
    db_module.db._conn = None

    async def noop(*args, **kwargs):
        return None

    def noop_sync(*args, **kwargs):
        return None

    async def task_processor_stub():
        try:
            while True:
                await asyncio.sleep(60)
        except asyncio.CancelledError:
            raise

    monkeypatch.setattr('app.main.ensure_collection', noop)
    monkeypatch.setattr('app.main.start_watcher', noop_sync)
    monkeypatch.setattr('app.main.stop_watcher', noop_sync)
    monkeypatch.setattr('app.main.run_task_processor', task_processor_stub)
    from app.main import app
    with TestClient(app) as test_client:
        yield test_client
    if db_module.db._conn is not None:
        asyncio.run(db_module.db.close())


@pytest.fixture(autouse=True)
def clear_drafts():
    documents_api._drafts.clear()
    yield
    documents_api._drafts.clear()


def test_edit_docx_replaces_body_table_and_header_without_mutating_original():
    original = make_docx()
    result = edit_docx(
        original,
        [
            {'find': '山田 太郎', 'replacement': '佐藤 花子'},
            {'find': '3,000円', 'replacement': '5,000円'},
            {'find': '2026年7月1日', 'replacement': '2026年7月28日'},
            {'find': '社内限定', 'replacement': '確認用ドラフト'},
        ],
    )
    assert result.total_replacements == 4
    assert '山田 太郎' in document_text(original)
    edited = document_text(result.content)
    assert '佐藤 花子' in edited
    assert '5,000円' in edited
    assert '2026年7月28日' in edited
    assert '確認用ドラフト' in edited


def test_edit_docx_rejects_invalid_or_no_match():
    with pytest.raises(ValueError, match='DOCX'):
        validate_docx_package(b'not-a-docx')
    with pytest.raises(ValueError, match='見つかりません'):
        edit_docx(make_docx(), [{'find': '存在しない', 'replacement': '値'}])


def test_edit_docx_preserves_surrounding_run_formatting():
    document = Document()
    paragraph = document.add_paragraph()
    prefix = paragraph.add_run('申請者: ')
    prefix.bold = True
    name_first = paragraph.add_run('山田 ')
    name_first.italic = True
    name_second = paragraph.add_run('太郎')
    suffix = paragraph.add_run('（本人）')
    suffix.underline = True
    output = io.BytesIO()
    document.save(output)

    result = edit_docx(
        output.getvalue(),
        [{'find': '山田 太郎', 'replacement': '佐藤 花子'}],
    )
    edited = Document(io.BytesIO(result.content)).paragraphs[0]
    assert edited.text == '申請者: 佐藤 花子（本人）'
    assert edited.runs[0].bold is True
    assert edited.runs[1].italic is True
    assert edited.runs[-1].underline is True


def test_docx_api_requires_preview_and_explicit_confirmation(client):
    headers = authenticated_headers(client)
    response = client.post(
        '/api/documents/docx/drafts',
        headers=headers,
        files={
            'file': (
                'expense.docx', make_docx(),
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
        },
        data={
            'operations_json': json.dumps([
                {'find': '山田 太郎', 'replacement': '佐藤 花子'},
                {'find': '3,000円', 'replacement': '5,000円'},
            ], ensure_ascii=False)
        },
    )
    assert response.status_code == 200
    preview = response.json()
    assert preview['total_replacements'] == 2
    assert 'まだ保存されていません' in preview['notice']
    draft_id = preview['draft_id']

    denied = client.post(
        f'/api/documents/docx/drafts/{draft_id}/confirm',
        headers=headers,
        json={'confirmed': False},
    )
    assert denied.status_code == 400

    saved = client.post(
        f'/api/documents/docx/drafts/{draft_id}/confirm',
        headers=headers,
        json={'confirmed': True},
    )
    assert saved.status_code == 200
    assert saved.headers['x-document-status'] == 'user-confirmed-draft'
    assert '佐藤 花子' in document_text(saved.content)
    assert '5,000円' in document_text(saved.content)

    second_download = client.post(
        f'/api/documents/docx/drafts/{draft_id}/confirm',
        headers=headers,
        json={'confirmed': True},
    )
    assert second_download.status_code == 404


def test_docx_api_rejects_non_docx(client):
    headers = authenticated_headers(client)
    response = client.post(
        '/api/documents/docx/drafts',
        headers=headers,
        files={'file': ('note.txt', b'hello', 'text/plain')},
        data={'operations_json': '[]'},
    )
    assert response.status_code == 400
